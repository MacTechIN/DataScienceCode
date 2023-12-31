import requests
from bs4 import BeautifulSoup
import time
import pyautogui
from docx import Document

# 사용자 입력
keyword = pyautogui.prompt("검색어를 입력하세요>>>")
lastpage = int(pyautogui.prompt("몇 페이지까지 크롤링 할까요?"))

# 워드 문서 생성하기
document = Document()

page_num = 1
for i in range(1, lastpage * 10, 10):
    print()
    print(f"{page_num}페이지 입니다. ============================================================================")
    print()
    response = requests.get(f"https://search.naver.com/search.naver?nso=&page={page_num}&qdt=-1&query={keyword}&qvt=-1&sm=tab_pag&start={i}&where=news")
    html = response.text
    soup = BeautifulSoup(html,'html.parser')
    links = soup.select("div.info_group") # 뉴스 기사 div 10개 추출

    for link in links:
        title = link.select("a.info") # 리스트
        if len(title) >= 2: # 링크가 2개 이상이면
            url = title[1].attrs['href'] # href의 속성값을 가져온다.
            response = requests.get(url, headers={'User-agent':'Mozila/5.0'})
            html = response.text
            soup = BeautifulSoup(html, 'html.parser')
            # print(url)
            # 만약 연예 뉴스 라면
            if "entertain" in response.url:
                title = soup.select_one(".end_tit")
                content = soup.select_one("#articeBody")
            elif "sports" in response.url:
                title = soup.select_one(".title")
                content = soup.select_one("#newsEndContents")
                #본문 내용안에 불필요한 div, p 삭제
                divs = content.select("div")
                for div in divs:
                    div.decompose()
                paragraphs = content.select("p")
                for p in paragraphs:
                    p.decompose()
            else:
                title = soup.select_one("#title_area")
                content = soup.select_one("#dic_area")
            # print(content.text)
            
            print("=======================링크=====================\n", url)
            print()
            print("=======================제목=====================\n", title.text.strip())
            print()
            print("=======================본문=====================\n", content.text.strip())
            
            # 워드에 링크, 제목, 본문 저장하기
            document.add_heading(title.text.strip())
            document.add_paragraph(url)
            document.add_paragraph(content.text.strip())

            time.sleep(0.3654654646416416541654165)
    page_num = page_num + 1  
    
# 워드 문서 저장하기
document.save(f"{keyword}_result.docx")   